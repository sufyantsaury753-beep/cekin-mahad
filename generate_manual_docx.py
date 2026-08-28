import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_callout(doc, title, text, bg_hex="F0FDF4", border_hex="059669"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, 140, 140, 180, 180)
    
    # Left border styling
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    
    r_text = p.add_run(text)
    r_text.font.name = "Arial"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_word_document():
    doc = docx.Document()
    
    # Page Setup (A4, standard margins)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Base Styling
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # =========================================================================
    # COVER / HEADER BANNER
    # =========================================================================
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_cell = header_table.rows[0].cells[0]
    h_cell.width = Inches(6.67)
    set_cell_background(h_cell, "064E3B") # Emerald 900
    set_cell_margins(h_cell, 240, 240, 240, 240)
    
    hp = h_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hr1 = hp.add_run("BUKU PANDUAN PENGGUNAAN SISTEM RESMI\n")
    hr1.font.name = "Arial"
    hr1.font.size = Pt(11)
    hr1.font.bold = True
    hr1.font.color.rgb = RGBColor(212, 175, 55) # Gold
    
    hr2 = hp.add_run("E-CHECKIN & PEMILIHAN KAMAR MA'HAD AL-JAMI'AH\n")
    hr2.font.name = "Arial"
    hr2.font.size = Pt(16)
    hr2.font.bold = True
    hr2.font.color.rgb = RGBColor(255, 255, 255)
    
    hr3 = hp.add_run("UIN SIBER SYEKH NURJATI CIREBON\n")
    hr3.font.name = "Arial"
    hr3.font.size = Pt(13)
    hr3.font.bold = True
    hr3.font.color.rgb = RGBColor(255, 255, 255)
    
    hr4 = hp.add_run("Portal Layanan: https://cekin-mahad.vercel.app | Tahun Ajaran 2026/2027")
    hr4.font.name = "Arial"
    hr4.font.size = Pt(9.5)
    hr4.font.color.rgb = RGBColor(167, 243, 208)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Metadata Box
    meta_tbl = doc.add_table(rows=2, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.rows[0].cells[0].paragraphs[0].add_run("🏢 Instansi: UPT Ma'had Al-Jami'ah UIN SSC\n📜 Dasar: SK Mudir No. B-092/Un.30/P.IV/2026").bold = False
    meta_tbl.rows[0].cells[1].paragraphs[0].add_run("📅 Jadwal Cek In: 19 - 21 Agustus 2026\n⏰ Waktu Pelayanan: 08.00 - 16.00 WIB").bold = False
    meta_tbl.rows[1].cells[0].paragraphs[0].add_run("📍 Asrama Putra: Gedung Ma'had Jadid (Lt. 2-5)").bold = False
    meta_tbl.rows[1].cells[1].paragraphs[0].add_run("📍 Asrama Putri: Gedung Qodim & Jadid (Lt. 2-5)").bold = False
    
    for row in meta_tbl.rows:
        for cell in row.cells:
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, 80, 80, 100, 100)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(71, 85, 105)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # =========================================================================
    # BAB 1: PENGENALAN SISTEM & STRUKTUR ASRAMA
    # =========================================================================
    p = doc.add_paragraph()
    r = p.add_run("BAB 1: PENGENALAN SISTEM & STRUKTUR ASRAMA")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(6, 78, 59)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph(
        "Sistem E-Checkin Ma'had Al-Jami'ah UIN Siber Syekh Nurjati Cirebon adalah aplikasi berbasis web modern "
        "yang mengotomatisasi seluruh alur prapendaftaran, pemilihan kamar asrama secara visual denah, penerbitan "
        "E-Tiket Boarding Pass dengan QR Code terenkripsi, hingga pemindaian check-in cepat oleh Mudabbir lorong pada Hari-H."
    )
    
    doc.add_paragraph("Pembagian 2 Gedung Asrama & Batasan Gender (L/P):").runs[0].bold = True
    
    gedung_tbl = doc.add_table(rows=3, cols=3)
    gedung_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Gedung Asrama", "Lantai Operasional", "Alokasi Gender & Nomor Kamar"]
    for i, h in enumerate(headers):
        cell = gedung_tbl.rows[0].cells[i]
        set_cell_background(cell, "065F46")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    data_g = [
        ("Gedung Ma'had Jadid", "Lantai 2 s/d 5", "👦 Putra (L): Kamar 09 - 16 (8 kamar per lantai)\n🧕 Putri (P): Kamar 01 - 08 & 17 - 24 (16 kamar per lantai)"),
        ("Gedung Ma'had Qodim", "Lantai 2 s/d 5", "🧕 Full Putri (P): Kamar 01 - 24 (24 kamar per lantai)")
    ]
    for row_idx, (g, l, d) in enumerate(data_g, start=1):
        r_cells = gedung_tbl.rows[row_idx].cells
        r_cells[0].paragraphs[0].add_run(g).bold = True
        r_cells[1].paragraphs[0].add_run(l)
        r_cells[2].paragraphs[0].add_run(d)
        for cell in r_cells:
            set_cell_background(cell, "FFFFFF" if row_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, 80, 80, 100, 100)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    add_callout(
        doc,
        "ATURAN PEMILIHAN KAMAR BERDASARKAN GENDER",
        "• Mahasantri Laki-laki (L) HANYA BISA memilih kamar di Gedung Jadid bagian Putra (Kamar 09 s/d 16).\n"
        "• Mahasantri Perempuan (P) HANYA BISA memilih kamar di Gedung Qodim (01-24) atau Gedung Jadid bagian Putri (01-08 & 17-24).\n"
        "• Sistem secara visual mengunci dan melarang pemilihan kamar lintas gender untuk menjaga privasi syar'i."
    )

    # =========================================================================
    # BAB 2: PANDUAN UNTUK MAHASANTRI
    # =========================================================================
    p = doc.add_paragraph()
    r = p.add_run("BAB 2: PANDUAN LANGKAH MAHASANTRI")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(6, 78, 59)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    
    steps_mhs = [
        ("Langkah 1: Verifikasi SK & Aktivasi PIN Keamanan",
         "1. Buka link https://cekin-mahad.vercel.app melalui HP atau Laptop.\n"
         "2. Klik menu 'Formulir Kamar' atau tombol 'Mulai Pemilihan Kamar'.\n"
         "3. Masukkan NISN (untuk Calon Mahasantri Baru) atau NIM (untuk Santri Lama/Perpanjangan).\n"
         "4. Klik tombol 'Cek Status di SK'.\n"
         "5. Jika terdaftar, buat 6 Digit PIN Rahasia (misal: 123456) dan masukkan No. WhatsApp aktif Anda.\n"
         "6. PIN ini berfungsi mengunci akun Anda agar data kamar tidak bisa diubah orang lain."),
        ("Langkah 2: Lengkapi Data & Unggah Pas Foto",
         "1. Data nama, jurusan, dan fakultas terkunci otomatis sesuai SK resmi Ma'had.\n"
         "2. Masukkan Nama Orang Tua/Wali dan No. WhatsApp Wali.\n"
         "3. Unggah Pas Foto Diri formal/rapi (bisa langsung jepret kamera HP atau pilih file galeri).\n"
         "4. Foto ini akan otomatis dicetak pada E-Tiket Boarding Pass Anda.\n"
         "5. Klik tombol 'Lanjut ke Pemilihan Kamar'."),
        ("Langkah 3: Memilih Gedung, Lantai, Kamar & Ranjang (Bed)",
         "1. Pilih Tab Gedung ('Ma'had Jadid' atau 'Ma'had Qodim').\n"
         "2. Pilih Lantai (Lantai 2, 3, 4, atau 5).\n"
         "3. Perhatikan denah kamar. Kotak kamar hijau menandakan kamar tersedia.\n"
         "4. Klik nomor kamar yang diinginkan, lalu pilih slot ranjang (Bed 1, 2, 3, atau 4) yang masih hijau.\n"
         "5. Klik tombol 'Konfirmasi & Terbitkan Tiket'."),
        ("Langkah 4: Unduh / Cetak E-Tiket Barcode",
         "1. E-Tiket Digital (Boarding Pass) Anda langsung terbit dengan Kode Barcode QR.\n"
         "2. Klik 'Cetak / Simpan PDF' atau simpan tangkapan layar (screenshot) tiket ke galeri HP Anda.\n"
         "3. Tiket ini WAJIB ditunjukkan ke Mudabbir lorong pada Hari-H kedatangan untuk scan check-in.")
    ]
    
    for title, desc in steps_mhs:
        p = doc.add_paragraph()
        r = p.add_run(f"▶ {title}")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        
        dp = doc.add_paragraph(desc)
        dp.paragraph_format.left_indent = Inches(0.2)
        dp.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # BAB 3: PANDUAN UNTUK MUDABBIR / PENGURUS LORONG
    # =========================================================================
    p = doc.add_paragraph()
    r = p.add_run("BAB 3: PANDUAN MUDABBIR & PENGURUS (SCANNER CHECK-IN)")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(6, 78, 59)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph(
        "Mudabbir (Putra) dan Mudabbirah (Putri) bertugas di setiap lantai lorong asrama pada tanggal 19–21 Agustus 2026 "
        "untuk memverifikasi kedatangan mahasantri dan menyerahkan kunci kamar."
    )
    
    steps_pgr = [
        ("1. Login Pengurus / Mudabbir",
         "• Buka menu 'Login' di pojok kanan atas website.\n"
         "• Pilih Tab 'Pengurus / Mudabbir'.\n"
         "• Masukkan Username/ID dan Password Anda (Default: mahad2026).\n"
         "• Anda akan otomatis diarahkan ke halaman Scanner Lorong."),
        ("2. Menjalankan Scanner Barcode Kamera",
         "• Klik menu 'Scanner Lorong'.\n"
         "• Izinkan peramban (browser) mengakses kamera HP Anda.\n"
         "• Arahkan lensa kamera ke Kode QR di E-Tiket mahasantri.\n"
         "• Jika kamera belakang belum aktif, klik tombol 'Ganti Kamera (Depan/Belakang)'.\n"
         "• Anda juga dapat menggunakan opsi 'Upload Foto Tiket' jika santri mengirim tiket via WhatsApp."),
        ("3. Verifikasi Data & Serah Terima Kunci",
         "• Saat QR terdeteksi, modal verifikasi mahasantri akan terbuka otomatis menampilkan: Nama, Pas Foto, Gedung, Lantai, Kamar, dan Nomor Bed.\n"
         "• Periksa apakah santri membawa barang sesuai tata tertib Ma'had.\n"
         "• Tulis catatan barang (opsional) pada form.\n"
         "• Klik tombol 'Konfirmasi Check-In & Berikan Kunci Kamar'.\n"
         "• Status santri seketika berubah menjadi CHECKED-IN dan tercatat di riwayat pengurus.")
    ]
    
    for title, desc in steps_pgr:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        
        dp = doc.add_paragraph(desc)
        dp.paragraph_format.left_indent = Inches(0.2)
        dp.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # BAB 4: PANDUAN SUPERADMIN MA'HAD
    # =========================================================================
    p = doc.add_paragraph()
    r = p.add_run("BAB 4: PANDUAN KONTROL SUPERADMIN MA'HAD")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(6, 78, 59)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    
    admin_features = [
        ("1. Tab Monitoring Check-In Santri",
         "• Memantau daftar seluruh mahasantri yang sudah memilih kamar maupun yang sudah check-in pada Hari-H.\n"
         "• Filter berdasarkan Gedung (Jadid/Qodim), Lantai (2-5), dan Status Kedatangan.\n"
         "• Tombol 'Ekspor CSV' di pojok kanan atas untuk mengunduh rekap laporan lengkap ke format Microsoft Excel."),
        ("2. Tab Manajemen Kamar & Kasur (+ / - Beds)",
         "• Fitur Kunci Kamar: Klik 'Kunci Kamar' pada nomor kamar tertentu (misal: Kamar Mudabbir/Kamar Tamu) agar tidak bisa dipilih oleh santri umum.\n"
         "• Fitur Ubah Kasur: Klik 'Ubah Kasur' untuk menambah atau mengurangi kapasitas kamar (misal: dari 4 kasur menjadi 6 kasur di lantai tertentu)."),
        ("3. Tab Master SK Mahasantri & Paginasi (800+ Data)",
         "• Menampilkan seluruh 800+ data santri resmi dari SK Ma'had.\n"
         "• Dilengkapi kontrol 'Tampilkan 50 / 100 / 250 / 500 / Semua Data' dan tombol halaman 1, 2, 3, dst.\n"
         "• Tombol 'Upload SK (PDF)' untuk mengimpor berkas lampiran SK baru secara massal."),
        ("4. Fitur Helpdesk (Reset PIN Santri)",
         "• Jika ada mahasantri yang lupa 6-digit PIN akunnya, Admin cukup mencari nama/NISN santri tersebut.\n"
         "• Klik tombol merah 'Reset PIN' pada kolom Aksi Helpdesk.\n"
         "• Akun santri akan kembali terbuka dan santri dapat mendaftar/membuat PIN baru."),
        ("5. Tab SK Mudabbir & Pengurus",
         "• Mengelola nama-nama Mudabbir/Mudabbirah resmi, penetapan lantai tugas, penetapan nomor kamar khusus mudabbir, dan reset password pengurus.")
    ]
    
    for title, desc in admin_features:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        
        dp = doc.add_paragraph(desc)
        dp.paragraph_format.left_indent = Inches(0.2)
        dp.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # BAB 5: FAQ & TROUBLESHOOTING
    # =========================================================================
    p = doc.add_paragraph()
    r = p.add_run("BAB 5: TANYA JAWAB (FAQ) & PENYELESAIAN MASALAH")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(6, 78, 59)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    
    faqs = [
        ("T: Bagaimana jika kamera scanner di HP pengurus terlihat gelap/tidak menyala?",
         "J: Pastikan izin akses kamera sudah diizinkan (Allow Camera Permission) di browser Chrome/Safari Anda. Anda juga dapat menekan tombol 'Ganti Kamera' atau menggunakan opsi 'Upload Foto Barcode'."),
        ("T: Bisakah mahasantri putra memilih kamar di Gedung Qodim?",
         "J: Tidak bisa. Gedung Ma'had Qodim dikhususkan 100% untuk mahasantri putri (P). Mahasantri putra (L) hanya dialokasikan di Gedung Ma'had Jadid kamar 09 s/d 16."),
        ("T: Bagaimana jika santri kehilangan file tiket cetaknya?",
         "J: Santri dapat login kembali kapan saja dengan NISN/NIM dan 6-Digit PIN rahasia mereka di website, lalu E-Tiket dapat diunduh ulang.")
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        r_q = p.add_run(f"❓ {q}\n")
        r_q.bold = True
        r_q.font.size = Pt(9.5)
        r_q.font.color.rgb = RGBColor(15, 23, 42)
        
        r_a = p.add_run(a)
        r_a.font.size = Pt(9.5)
        r_a.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_after = Pt(6)

    # Save to file
    out_dir = r"C:\Users\sufya\.gemini\antigravity\scratch\cekin-mahad\public"
    os.makedirs(out_dir, exist_ok=True)
    out_path_public = os.path.join(out_dir, "Panduan_Penggunaan_Cekin_Mahad_UINSSC.docx")
    out_path_root = r"C:\Users\sufya\.gemini\antigravity\scratch\cekin-mahad\Panduan_Penggunaan_Cekin_Mahad_UINSSC.docx"
    
    doc.save(out_path_public)
    doc.save(out_path_root)
    print("SUCCESS: Document created at:", out_path_root)

if __name__ == "__main__":
    build_word_document()
